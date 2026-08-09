import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout(doc, text, title="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "EFF6FF")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="2563EB"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(37, 99, 235)
    
    r_text = p.add_run(text)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    r_text.font.size = Pt(10)
    
    doc.add_paragraph("")

def add_tech_entry(doc, name, version, category, purpose, why_used, key_features, benefits, alternatives, adoption):
    h = doc.add_heading(level=3)
    r = h.add_run(f"• {name}" + (f" ({version})" if version else ""))
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(30, 41, 59)
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    runs_data = [
        ("Category: ", category),
        ("\nPurpose: ", purpose),
        ("\nWhy Used in GarudaAI: ", why_used),
        ("\nKey Features: ", key_features),
        ("\nBenefits: ", benefits),
        ("\nAlternative Technologies: ", alternatives),
        ("\nIndustry Adoption: ", adoption)
    ]
    
    for label, val in runs_data:
        r_lbl = p.add_run(label)
        r_lbl.bold = True
        r_lbl.font.size = Pt(10)
        r_lbl.font.color.rgb = RGBColor(71, 85, 105)
        
        r_val = p.add_run(val)
        r_val.font.size = Pt(10)
        r_val.font.color.rgb = RGBColor(30, 41, 59)
        
    p_space = doc.add_paragraph("")
    p_space.paragraph_format.space_after = Pt(2)

def generate_document(output_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    PRIMARY = RGBColor(37, 99, 235)      # Royal Blue
    SECONDARY = RGBColor(15, 23, 42)    # Dark Slate
    MUTED = RGBColor(100, 116, 139)     # Slate Gray
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = SECONDARY
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("GarudaAI Security Intelligence Platform")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Comprehensive Technology Stack Specification & Architecture Guide")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = MUTED
    
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run("Version: 1.0.0  |  Classification: Technical Architecture Documentation  |  Status: Verified Stack")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = MUTED
    
    doc.add_paragraph("―" * 55)
    
    # 1. PROJECT OVERVIEW
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Project Overview")
    r1.font.color.rgb = PRIMARY
    r1.font.bold = True
    
    p = doc.add_paragraph(
        "GarudaAI is an enterprise-grade insider threat intelligence, zero-trust access management, and behavior trust analytics platform. "
        "Designed to prevent data exfiltration and credential misuse within high-security enterprise networks, GarudaAI dynamically correlates multi-source "
        "activity logs (authentication locations, data downloads, psychometric risk factors, time anomalies) into a real-time Behavior Trust Score (0–100).\n\n"
        "When anomalous risk spikes occur, GarudaAI automatically triggers an AI-powered incident investigation playbook driven by Google Gemini, "
        "provides Just-In-Time Privileged Access Management (JIT PAM) enforcement, and secures data streams using NIST FIPS 203 Post-Quantum Cryptography (ML-KEM-768)."
    )
    p.paragraph_format.line_spacing = 1.15
    
    # 2. TECHNOLOGY STACK SUMMARY
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Technology Stack Summary")
    r2.font.color.rgb = PRIMARY
    r2.font.bold = True
    
    doc.add_paragraph(
        "• Frontend: React 19 SPA built with Vite 8, styled using Tailwind CSS v3 with dark-mode glassmorphic aesthetics, Chart.js v4, and Lucide React icons.\n"
        "• Backend: Modular Python 3 Flask REST API (v3.0.3) with rate limiting (Flask-Limiter), CORS security (Flask-Cors), and schema validation (Pydantic v2).\n"
        "• AI & Machine Learning: Dual AI engine combining a Scikit-Learn Random Forest Classifier for quantitative threat scoring and Google Gemini AI (google-genai SDK >=1.0.0) for qualitative incident reports and SOAR playbooks.\n"
        "• Database: MongoDB (PyMongo v4.7.2 driver) managing schema-less documents for security logs, employee profiles, psychometrics, and JIT records.\n"
        "• Security & PQC: NIST FIPS 203 ML-KEM-768 via liboqs-python (with Python standard library fallback), HKDF-SHA256 derivation, and AES-256-GCM authenticated payload encryption.\n"
        "• Authentication: Firebase Authentication (SDK v12.16.0 frontend, firebase-admin v6.5.0 backend) paired with custom salted-hash RBAC database."
    )
    
    # 3. FRONTEND STACK
    h3_sec = doc.add_heading(level=1)
    r3_sec = h3_sec.add_run("3. Frontend Stack")
    r3_sec.font.color.rgb = PRIMARY
    r3_sec.font.bold = True
    
    add_tech_entry(
        doc, "React", "^19.2.7", "Frontend Framework / UI Library",
        "Provides a component-driven, declarative architecture for rendering interactive web dashboards and real-time security UI elements.",
        "React's virtual DOM and modular component structure enable fast re-rendering of live threat scores, search filters, modal dialogs, and activity timelines without full page reloads.",
        "React Hooks (useState, useEffect, useContext), Virtual DOM, JSX syntax, Component Lifecycle management.",
        "Excellent UI performance, modular reusable components, extensive ecosystem, robust developer tooling.",
        "Vue.js, Svelte, Angular.",
        "Industry standard for enterprise web application dashboards (used by Meta, Netflix, Airbnb, Microsoft)."
    )
    
    add_tech_entry(
        doc, "React DOM", "^19.2.7", "Web DOM Rendering Engine",
        "Serves as the entry point to the browser DOM for React components.",
        "Required by React to mount and manipulate virtual DOM trees into real web browser DOM nodes (index.html).",
        "Efficient DOM diffing, event delegation, client-side rendering.",
        "Seamless sync between React component states and DOM representations.",
        "Preact DOM.",
        "Universal requirement for React web applications."
    )
    
    add_tech_entry(
        doc, "Vite", "^8.1.1", "Build Tool & Development Server",
        "Provides hyper-fast Hot Module Replacement (HMR) during local development and bundles optimized production assets.",
        "Vite leverages native ES modules in modern browsers to instantly boot local development (npm run dev) and provides transparent API proxying (/api -> http://localhost:5000).",
        "Instant server start, lightning-fast HMR, Rollup-based production bundling, built-in API proxy configuration.",
        "Dramatically speeds up developer iteration cycles compared to legacy Webpack builds.",
        "Webpack, Turbopack, Parcel, esbuild.",
        "Modern industry standard for bundling React, Vue, and frontend applications."
    )
    
    add_tech_entry(
        doc, "Tailwind CSS", "^3.4.1", "Utility-First CSS Framework",
        "Provides utility classes for constructing custom, responsive, dark-mode glassmorphism UI layouts.",
        "Allows building high-density SOC analytics screens with glassmorphism backgrounds (backdrop-blur), subtle borders, status pills, and threat level colors without writing custom CSS files.",
        "JIT compiler, utility classes, theme customization (tailwind.config.js), dark mode variants.",
        "Accelerated design-to-code workflow, zero unused CSS in production bundles, precise layout control.",
        "Bootstrap, Chakra UI, Material-UI (MUI), Bulma.",
        "Leading utility CSS framework across modern full-stack web applications."
    )
    
    add_tech_entry(
        doc, "Chart.js", "^4.5.1", "Data Visualization / Charting Library",
        "Renders HTML5 Canvas-based interactive charts for tracking employee trust score trends over time.",
        "GarudaAI displays historical risk trajectories, trust score declines, and behavioral anomalies using interactive time-series line charts and gauge components.",
        "Canvas rendering, animation support, tooltips, responsive canvas sizing, multi-axis formatting.",
        "Smooth frame-rate animations, lightweight bundle footprint, highly customizable canvas rendering.",
        "Recharts, D3.js, ECharts, ApexCharts.",
        "Widely used across analytics platforms, financial dashboards, and SIEM monitoring tools."
    )
    
    add_tech_entry(
        doc, "Lucide React", "^1.24.0", "Icon Library",
        "Provides clean, scalable SVG vector icons for security navigation, alerts, status indicators, and modal buttons.",
        "Delivers consistent, security-oriented visual indicators (Shields, Key, Lock, AlertTriangle, Activity, User, Terminal) aligning with dark-mode SOC aesthetics.",
        "Tree-shakeable SVG components, customizable stroke width, color, and size props.",
        "Lightweight footprint, consistent design language across all dashboard components.",
        "FontAwesome, Heroicons, Material Icons, Feather Icons.",
        "Standard choice for modern React design systems (Shadcn/UI default)."
    )

    # 4. BACKEND STACK
    h4_sec = doc.add_heading(level=1)
    r4_sec = h4_sec.add_run("4. Backend Stack")
    r4_sec.font.color.rgb = PRIMARY
    r4_sec.font.bold = True
    
    add_tech_entry(
        doc, "Python", "3.11 / 3.12 / 3.14", "Core Backend Programming Language",
        "Executes API endpoints, threat scoring algorithms, cryptographic key encapsulation, and AI orchestration.",
        "Python is the native ecosystem for data science, machine learning (Scikit-Learn), cryptography (liboqs, cryptography), and official Google Gemini AI SDKs (google-genai).",
        "Strong typing extensions, rich standard libraries (hashlib, socket, ssl), broad AI/ML ecosystem.",
        "High readability, rapid security engineering prototyping, native ML library support.",
        "Node.js (TypeScript), Go, Rust, Java.",
        "Dominant programming language in AI/ML, Cyber Security, and SOC automation."
    )
    
    add_tech_entry(
        doc, "Flask", "==3.0.3", "Micro Web Backend Framework",
        "Serves RESTful HTTP API routes (/api/employees, /api/ai/explain, /api/pqc/*, /api/jit/*).",
        "Lightweight WSGI architecture allows modular route blueprint organization (jit_bp) and direct integration with Python security/ML algorithms without heavy boilerplate.",
        "WSGI application object, Request/Response abstractions, Blueprint modular routing, HTTP error handling.",
        "Minimal latency overhead, total control over middleware stack and API structures.",
        "FastAPI, Django, Express.js (Node.js).",
        "Extensively used in microservices, ML inference APIs, and security tools (Pinterest, Netflix, LinkedIn)."
    )
    
    add_tech_entry(
        doc, "Flask-CORS", "==4.0.1", "HTTP Middleware / CORS Security",
        "Enables and manages Cross-Origin Resource Sharing (CORS) headers for incoming API requests.",
        "Safely permits the Vite client running on localhost:5173 to exchange JSON payloads, authorization headers (X-JIT-Token, Authorization), and credentials with the Flask backend on localhost:5000.",
        "Origin matching, allowed headers specification, pre-flight OPTIONS request handling.",
        "Prevents browser security blocks while maintaining strict cross-origin access rules.",
        "Custom CORS header wrappers, Nginx reverse proxy CORS rules.",
        "Standard CORS extension for Flask applications."
    )
    
    add_tech_entry(
        doc, "Flask-Limiter", "==3.7.0", "API Rate Limiting Middleware",
        "Enforces rate limits on HTTP endpoints to protect against brute-force attacks and service exhaustion.",
        "GarudaAI enforces a default limit of 100 requests per minute per client IP (get_remote_address) to safeguard authentication and AI endpoint capacity.",
        "IP-based tracking, configurable rate limit strings, HTTP 429 Too Many Requests response handling.",
        "Prevents API abuse, protects Gemini quota consumption, and mitigates DDoS attempts.",
        "Redis-based custom rate limiters, Nginx rate-limiting directives.",
        "Standard rate-limiting extension in Python WSGI ecosystems."
    )
    
    add_tech_entry(
        doc, "Pydantic", ">=2.0.0", "Data Validation & Schema Enforcement",
        "Validates incoming request payloads and structures internal data transfer objects (DTOs).",
        "Ensures strictly-typed input data for user accounts, JIT access requests, and threat simulation parameters before processing.",
        "Type annotation validation, custom validators, JSON schema generation, high performance (Rust core).",
        "Prevents malformed data injection and eliminates runtime type errors.",
        "Cerberus, Marshmallow, Schematics.",
        "De-facto standard data validation engine in modern Python APIs (FastAPI core, Pydantic v2)."
    )

    # 5. AI / MACHINE LEARNING STACK
    h5_sec = doc.add_heading(level=1)
    r5_sec = h5_sec.add_run("5. Artificial Intelligence & Machine Learning Stack")
    r5_sec.font.color.rgb = PRIMARY
    r5_sec.font.bold = True
    
    add_tech_entry(
        doc, "Google GenAI SDK (google-genai)", ">=1.0.0", "Generative AI SDK & Cloud Client",
        "Connects the backend server to Google Gemini Large Language Models (gemini-2.5-flash, gemini-2.0-flash, gemini-1.5-flash).",
        "Powers GarudaAI's core intelligence: generating structured incident investigation reports, human-readable risk explanations, threat mitigation playbooks, and natural language database threat queries.",
        "Native Python SDK interface, streaming support, resilient socket connection validation, built-in fallback mechanisms.",
        "High-speed inference (Flash models), state-of-the-art cybersecurity reasoning, native integration with Google Cloud infrastructure.",
        "OpenAI API SDK, Anthropic Claude SDK, Local Ollama / Llama 3 models.",
        "Official Google Cloud SDK for Next-Gen Gemini AI integrations."
    )
    
    add_tech_entry(
        doc, "Scikit-Learn (scikit-learn)", "Latest", "Machine Learning Framework",
        "Trains and evaluates the core Random Forest Classifier (insider_threat_rf.joblib) for predicting insider threat probabilities.",
        "GarudaAI extracts behavioral metrics (after-hours logins, file download spikes, USB insertions, psychometric OCEAN traits) and feeds them to an ensemble Random Forest model to score risk.",
        "Ensemble learning (RandomForestClassifier), dataset partitioning (train_test_split), cross-validation (cross_val_score), classification metrics.",
        "Robust against overfitting, high interpretability via feature importance weights, fast evaluation times.",
        "XGBoost, LightGBM, TensorFlow, PyTorch.",
        "Gold standard Python library for traditional supervised machine learning and tabular security data classification."
    )
    
    add_tech_entry(
        doc, "Joblib", "Latest", "Model Persistence & Serialization",
        "Serializes trained Scikit-Learn Random Forest model binaries to disk (insider_threat_rf.joblib) and loads them into memory for instant inference.",
        "Avoids model re-training on every backend restart by providing fast binary serialization of the trained classifier model.",
        "Efficient NumPy array compression, fast disk read/write primitives.",
        "Zero latency overhead during API startup; instant load of pre-trained weights.",
        "Pickle, ONNX runtime, ONNX-ML.",
        "Default model persistence standard recommended by Scikit-Learn."
    )

    # 6. DATABASE STACK
    h6_sec = doc.add_heading(level=1)
    r6_sec = h6_sec.add_run("6. Database Stack")
    r6_sec.font.color.rgb = PRIMARY
    r6_sec.font.bold = True
    
    add_tech_entry(
        doc, "MongoDB", "Community / Atlas", "NoSQL Document Database",
        "Serves as the primary persistence layer for GarudaAI's unstructured log streams, employee documents, and system state.",
        "Security logs, psychometric traits, timeline event nodes, and JIT elevation records possess flexible, evolving schemas mapping naturally to JSON/BSON documents.",
        "BSON storage, dynamic schema design, flexible indexed queries, high insert throughput.",
        "Scales horizontally, handles diverse log formats without schema migrations, enables fast lookup by employee_id.",
        "PostgreSQL (with JSONB), Elasticsearch, Couchbase.",
        "Widely adopted for operational log monitoring, SIEM data lakes, and flexible modern web backends."
    )
    
    add_tech_entry(
        doc, "PyMongo", "==4.7.2", "Database Driver",
        "Provides official Python interface for establishing client connections (MongoClient) and running queries against MongoDB collections.",
        "Offers reliable connection pooling, direct BSON dictionary mapping, and native Python exception handling.",
        "Connection pooling, thread safety, indexing helpers, CRUD operations.",
        "Lightweight and battle-tested driver with high execution speed.",
        "MongoEngine, Motor (async driver).",
        "Official standard Python driver for MongoDB."
    )

    # 7. SECURITY & PQC STACK
    h7_sec = doc.add_heading(level=1)
    r7_sec = h7_sec.add_run("7. Security & Post-Quantum Cryptography Stack")
    r7_sec.font.color.rgb = PRIMARY
    r7_sec.font.bold = True
    
    add_tech_entry(
        doc, "Post-Quantum Cryptography (NIST FIPS 203 ML-KEM-768)", "liboqs >=0.15.0", "Post-Quantum Public Key Cryptography",
        "Implements Module-Lattice-Based Key Encapsulation Mechanism (ML-KEM-768 / Kyber768) to protect session key exchanges against future quantum decryption attacks.",
        "Establishes quantum-resistant secure channels for exchanging privilege tokens and sensitive threat intelligence logs.",
        "1184-byte Public Key, 2400-byte Private Key, 1088-byte Ciphertext, 32-byte Shared Secret. Dual-engine architecture guarantees native speed when available with standard math fallback.",
        "Future-proof quantum security compliance (NIST FIPS 203 standard).",
        "FrodoKEM, Classic McEliece, Falcon.",
        "Official NIST standard for post-quantum key encapsulation (selected by NSA, CISA, Google, Cloudflare)."
    )
    
    add_tech_entry(
        doc, "Python Cryptography (cryptography)", ">=46.0.0", "Symmetric Encryption Primitives",
        "Provides AES-256-GCM authenticated symmetric encryption for encrypting sensitive fields, session payloads, and local tokens.",
        "Combines 256-bit AES confidential encryption with Galois Counter Mode (GCM) integrity authentication tags (128-bit) to prevent tampering.",
        "AES-GCM cipher initialization, Nonce management (96-bit NIST standard), AEAD (Authenticated Encryption with Associated Data).",
        "Cryptographically secure against ciphertext forgery and eavesdropping.",
        "PyCryptodome, NaCl / libsodium.",
        "Industry gold standard Python cryptographic library maintained by PyCA."
    )
    
    add_tech_entry(
        doc, "HKDF (HMAC-based Key Derivation)", ">=0.0.3", "Cryptographic Key Derivation",
        "Expands shared secret bytes output by the PQC ML-KEM encapsulation process into cryptographically uniform 256-bit symmetric key material.",
        "Derives AES-256 keys using HKDF-SHA256 with custom application salts (GarudaAI-PQC-Salt-v1-2026) and domain-separated context tags.",
        "Extract-and-Expand workflow, SHA-256 pseudorandom function (PRF).",
        "Ensures maximum entropy for derived symmetric keys.",
        "PBKDF2, Argon2.",
        "RFC 5869 standard used in TLS 1.3, Signal protocol, and IPSec."
    )
    
    add_tech_entry(
        doc, "JIT Privileged Access Management Engine", "Custom", "Zero-Trust Access Control",
        "Enforces time-bound, ephemeral administrative access grants with automatic expiration timers and audit logging.",
        "Eliminates permanent standing admin privileges across corporate infrastructure by requiring approval-backed temporary access windows.",
        "Ephemeral token generation, automatic expiration checks, status tracking (APPROVED, EXPIRED, DENIED), audit trail logging.",
        "Reduces attack surface and mitigates stolen admin credential abuse.",
        "HashiCorp Vault SSH JIT, Teleport, CyberArk PAM.",
        "Core pillar of modern Zero Trust Architecture (ZTA / NIST SP 800-207)."
    )

    # 8. AUTHENTICATION & DEVOPS
    h8_sec = doc.add_heading(level=1)
    r8_sec = h8_sec.add_run("8. Authentication, DevOps & Deployment Stack")
    r8_sec.font.color.rgb = PRIMARY
    r8_sec.font.bold = True
    
    add_tech_entry(
        doc, "Firebase Authentication", "SDK v12.16.0 / Admin v6.5.0", "Identity Management Service",
        "Handles user sign-in, password reset emails, JWT token generation, and secure session management.",
        "Offloads identity management to Google Cloud infrastructure while offering graceful fallback to local development RBAC credentials when offline.",
        "Email/Password authentication, ID token verification, session state observers (onAuthStateChanged).",
        "Scalable identity platform with built-in OAuth providers and security monitoring.",
        "Auth0, AWS Cognito, Keycloak, Supabase Auth.",
        "Global cloud identity standard for web and mobile applications."
    )
    
    add_tech_entry(
        doc, "Vercel", "Production", "Cloud Application Hosting / Frontend CDN",
        "Hosts the built Vite React single-page application on global edge networks.",
        "Provides single-command deployment, dynamic route rewrites, SSL provisioning, and continuous deployment triggers.",
        "Edge network distribution, automated CI/CD builds, client-side SPA routing support.",
        "Low latency access worldwide with zero server infrastructure maintenance.",
        "Netlify, AWS CloudFront + S3, Cloudflare Pages.",
        "Primary cloud deployment target for modern React/Vite frontends."
    )
    
    add_tech_entry(
        doc, "Render / Local WSGI", "Production", "Backend Application Hosting Environment",
        "Runs the Python Flask WSGI API server and manages environment secret injection.",
        "Supports continuous deployment directly from Git repositories, automatically installs requirements.txt, and manages environment key bindings.",
        "Managed container execution, HTTPS endpoints, automatic git push deployment.",
        "Simplifies backend infrastructure setup without manual VM maintenance.",
        "AWS Elastic Beanstalk, Heroku, DigitalOcean App Platform.",
        "Popular platform for hosting Python microservices and web APIs."
    )

    # 9. TESTING & CODE QUALITY
    h9_sec = doc.add_heading(level=1)
    r9_sec = h9_sec.add_run("9. Testing & Code Quality Tools")
    r9_sec.font.color.rgb = PRIMARY
    r9_sec.font.bold = True
    
    add_tech_entry(
        doc, "Pytest", "Latest", "Python Test Framework",
        "Executes automated unit and integration tests across security, database, and AI modules (test_pqc.py, test_auth_db.py, test_jit_pam.py, test_gemini_integration.py, test_backend.py).",
        "Guarantees code stability and security contract validation prior to production releases.",
        "Test fixture injection, assertion introspection, detailed failure output.",
        "High test execution speed, fixture reusability, rich plugin ecosystem.",
        "Unittest, Robot Framework.",
        "De-facto Python standard testing framework."
    )
    
    add_tech_entry(
        doc, "Oxlint", "^1.71.0", "JavaScript / React Code Linter",
        "Performs static analysis on React JSX source files to catch syntax errors and performance antipatterns.",
        "Written in Rust, Oxlint is up to 50x faster than traditional ESLint, offering instant feedback during development (npm run lint).",
        "High-speed static analysis, built-in React JSX rule sets.",
        "Enforces consistent frontend code quality with minimal build overhead.",
        "ESLint, Biome.",
        "Rapidly growing adoption in modern web development stacks."
    )

    # 10. ARCHITECTURE FLOW & INTERACTIVITY
    h10_sec = doc.add_heading(level=1)
    r10_sec = h10_sec.add_run("10. Architecture Flow & End-to-End Trace")
    r10_sec.font.color.rgb = PRIMARY
    r10_sec.font.bold = True
    
    p_arch = doc.add_paragraph(
        "1. Authentication: The user opens the React frontend. Firebase Auth (or local RBAC fallback) authenticates credentials.\n"
        "2. Behavioral Ingestion: React sends HTTP GET requests via fetch to Flask endpoints (/api/employees), passing CORS and rate limit checks.\n"
        "3. ML Risk Calculation: Flask passes behavioral metrics to the Scikit-Learn Random Forest Engine, computing threat probabilities and updating the Behavior Trust Score (0–100) in MongoDB.\n"
        "4. Post-Quantum Encryption: Elevated operations trigger the PQC ML-KEM-768 engine, generating keypairs, deriving symmetric keys via HKDF-SHA256, and encrypting payloads using AES-256-GCM.\n"
        "5. Zero-Trust Access: Temporary admin requests pass through the JIT PAM engine (jit_middleware.py), enforcing time-bound expiration timers.\n"
        "6. Gemini AI Playbook Generation: When trust scores drop, backend calls gemini_service (google-genai SDK), synthesizing natural language threat reports and SOAR playbooks.\n"
        "7. Dashboard Visualization: React client receives JSON payloads, rendering historical risk trends using Chart.js and badging indicators via Tailwind CSS."
    )
    p_arch.paragraph_format.line_spacing = 1.15

    # 11. TECHNOLOGY SUMMARY TABLE
    h11_sec = doc.add_heading(level=1)
    r11_sec = h11_sec.add_run("11. Technology Summary Table")
    r11_sec.font.color.rgb = PRIMARY
    r11_sec.font.bold = True
    
    summary_data = [
        ("Frontend", "React", "^19.2.7", "Renders reactive, modular SOC analytics dashboards"),
        ("Frontend", "Vite", "^8.1.1", "Enables instant HMR dev server & Rollup bundling"),
        ("Frontend", "Tailwind CSS", "^3.4.1", "Builds dark-mode glassmorphic security UI fast"),
        ("Frontend", "Chart.js", "^4.5.1", "Visualizes user trust score trends over time"),
        ("Frontend", "Lucide React", "^1.24.0", "Delivers crisp vector icons for cybersecurity states"),
        ("Backend", "Python", "3.11+", "Executes backend REST API, security, and AI logic"),
        ("Backend", "Flask", "==3.0.3", "Provides modular, lightweight RESTful API endpoints"),
        ("Backend", "Flask-CORS", "==4.0.1", "Enables safe cross-origin API calls (5173 -> 5000)"),
        ("Backend", "Flask-Limiter", "==3.7.0", "Protects API endpoints against DDoS & rate abuse"),
        ("Backend", "Pydantic", ">=2.0.0", "Validates request structures and prevents bad data"),
        ("AI / ML", "Google GenAI SDK", ">=1.0.0", "Generates automated threat reports & SOAR playbooks"),
        ("AI / ML", "Scikit-Learn", "Latest", "Predicts insider threat risk using Random Forest"),
        ("AI / ML", "Joblib", "Latest", "Loads trained Random Forest binary weights instantly"),
        ("AI / ML", "NumPy", "Latest", "Executes matrix operations for feature extraction"),
        ("AI / ML", "Matplotlib", "Latest", "Generates ROC and confusion matrix evaluation images"),
        ("Database", "MongoDB", "Latest", "Stores security logs, user profiles, and JIT records"),
        ("Database", "PyMongo", "==4.7.2", "Connects Flask backend cleanly to MongoDB"),
        ("Security", "NIST ML-KEM-768", ">=0.15.0", "Safeguards key exchanges against quantum attacks"),
        ("Security", "Cryptography", ">=46.0.0", "Provides AES-256-GCM authenticated encryption"),
        ("Security", "HKDF", ">=0.0.3", "Derives symmetric AES keys from raw PQC secrets"),
        ("Security", "JIT PAM Engine", "Custom", "Limits high-risk admin access to temporary windows"),
        ("Auth", "Firebase Auth", "SDK/Admin", "Manages authentication with local RBAC fallback"),
        ("DevOps", "Vercel", "Production", "Distributes React application globally via CDN"),
        ("DevOps", "Render", "Production", "Hosts Flask backend and manages environment keys"),
        ("Testing", "Pytest", "Latest", "Validates security, PQC, JIT, and API contracts"),
        ("Quality", "Oxlint", "^1.71.0", "Enforces frontend code quality and catches errors"),
        ("Automation", "Windows Batch", "Custom", "Launches full platform locally with single click")
    ]
    
    tbl = doc.add_table(rows=len(summary_data) + 1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    headers = ["Category", "Technology", "Version", "Why Used in GarudaAI"]
    hdr_cells = tbl.rows[0].cells
    col_widths = [Inches(1.2), Inches(1.5), Inches(1.0), Inches(2.8)]
    
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "2563EB")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.runs[0]
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    for row_idx, data in enumerate(summary_data):
        row_cells = tbl.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=100, right=100)
            p = row_cells[col_idx].paragraphs[0]
            if len(p.runs) > 0:
                p.runs[0].font.size = Pt(9.0)
                p.runs[0].font.color.rgb = SECONDARY
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F8FAFC")
                
    doc.add_paragraph("")

    # 12. FUTURE ROADMAP
    h12_sec = doc.add_heading(level=1)
    r12_sec = h12_sec.add_run("12. Suggested Future Enhancements")
    r12_sec.font.color.rgb = PRIMARY
    r12_sec.font.bold = True
    
    add_callout(
        doc,
        "The technologies listed below are not currently part of the active GarudaAI codebase. "
        "They represent strategic recommendations for future enterprise scaling and production hardening.",
        title="SUGGESTED FUTURE ENHANCEMENTS"
    )
    
    p_fut = doc.add_paragraph(
        "1. Redis Caching Layer: Introduce Redis for distributed API rate limiting (replacing in-memory Flask-Limiter) and real-time session token revocation caching.\n"
        "2. Apache Kafka Stream Processing: Integrate Apache Kafka for real-time ingestion and pipeline parsing of multi-terabyte corporate SIEM log streams.\n"
        "3. OpenTelemetry Audit Tracing: Implement OpenTelemetry APM agents across Flask and MongoDB for distributed request tracing and performance bottleneck monitoring.\n"
        "4. PostgreSQL / TimescaleDB Time-Series Data Lake: Partition historical security log metrics into a dedicated time-series database to optimize long-term analytical queries.\n"
        "5. Docker & Kubernetes (K8s) Orchestration: Package services into OCI-compliant container images managed via Kubernetes manifests for automated auto-scaling and zero-downtime rolling updates."
    )
    p_fut.paragraph_format.line_spacing = 1.15

    doc.save(output_path)
    print(f"Word Document successfully saved to: {output_path}")

if __name__ == "__main__":
    # Use the project root directory (two levels up from scripts/)
    out_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    out_file = os.path.join(out_dir, "GarudaAI_Tech_Stack_Specification.docx")
    generate_document(out_file)
