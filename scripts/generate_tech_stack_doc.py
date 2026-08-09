import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_doc():
    doc = Document()
    
    # Page margins
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    PRIMARY_COLOR = RGBColor(37, 99, 235)  # Blue Accent
    DARK_TEXT = RGBColor(15, 23, 42)       # Slate 900
    GRAY_TEXT = RGBColor(100, 116, 139)    # Slate 500
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = DARK_TEXT
    
    # Title
    title = doc.add_paragraph()
    run = title.add_run("GarudaAI Threat Intelligence Platform")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR
    
    # Subtitle
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Technology Stack, System Architecture & Machine Learning Guide")
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = GRAY_TEXT
    
    doc.add_paragraph("_" * 60)
    doc.add_paragraph("")
    
    # SECTION 1: Technology Stack
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Unified Technology Stack")
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph(
        "GarudaAI is architected as a decoupled web application containing a React client, a Flask API server, "
        "and dual-mode database layers (MongoDB for production and local JSON for fallback demo deployment). "
        "The model layer combines rule-based indicators with classical machine learning (Random Forest) "
        "and generative AI playbooks (Google Gemini)."
    )
    
    # Table of Tech Stack
    table = doc.add_table(rows=6, cols=3)
    table.autofit = False
    
    headers = ["Component", "Technology", "Description & Role"]
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_background(hdr_cells[i], "2563EB")
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
    stack_data = [
        ("Frontend View", "React 19, Vite, Chart.js", "Dashboard UI rendering chronological user timeline charts."),
        ("Backend Server", "Python 3, Flask, CORS", "API controller serving score models and database logs."),
        ("Database Tier", "MongoDB / Mock DB wrapper", "Collections caching raw activity files (logon, file, device, http, email)."),
        ("AI Playbooks", "Google Gemini API", "Synthesizes security incident summaries and SOAR playbooks."),
        ("ML Evaluation", "scikit-learn, joblib", "Random Forest Classifier predicting insider threat profiles.")
    ]
    
    for row_idx, data in enumerate(stack_data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_margins(row_cells[col_idx], top=100, bottom=100)
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F8FAFC")
                
    doc.add_paragraph("")
    
    # SECTION 2: Architecture & Folder Structure
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. System Folder Architecture")
    h2_run.font.bold = True
    h2_run.font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph("The structure below maps the key folders and components of the workspace:")
    
    # Code block display
    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.5)
    code_run = code_p.add_run(
        "FINSpark/\n"
        "├── evaluate_model.py          # ML Model training & evaluation pipeline\n"
        "├── insider_threat_rf.joblib   # Serialized trained Random Forest model\n"
        "├── backend/                   # Flask REST API Server\n"
        "│   ├── app.py                 # Core server routes & Flask settings\n"
        "│   ├── trust_score.py         # Algorithmic Trust Score logic\n"
        "│   ├── timeline.py            # Collapses routine events for display\n"
        "│   └── db_client.py           # MongoDB Client + fallback JSON database\n"
        "├── frontend/                  # React Client Side\n"
        "│   ├── src/\n"
        "│   │   ├── App.jsx            # Core UI layout components\n"
        "│   │   └── components/        # Isolated JIT PAM & Login views\n"
        "└── dataset/                   # Carnegie Mellon CERT v4.2 dataset\n"
        "    ├── employees.csv          # Metadata for all 1000 employees\n"
        "    └── ground_truth_labels.csv# Insider threat target labels (150 rows)\n"
    )
    code_run.font.name = 'Consolas'
    code_run.font.size = Pt(9.5)
    code_run.font.color.rgb = GRAY_TEXT
    
    # SECTION 3: Threat Detection Pipeline
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Dual-Model Threat Detection Pipeline")
    h3_run.font.bold = True
    h3_run.font.color.rgb = PRIMARY_COLOR
    
    doc.add_paragraph(
        "GarudaAI operates a dual-model pipeline. It tracks daily behaviors deterministically "
        "using a Behavior Trust Scoring Engine and predicts profiles using a Machine Learning Classifier."
    )
    
    doc.add_heading("3.1 Behavior Trust scoring Engine", level=2)
    doc.add_paragraph(
        "Every user begins with a base score of 100. Deductions are subtracted based on severity violations:\n"
        "• Off-Hours Authentication: -5.0 pts\n"
        "• Unrecognized Terminal Logon: -10.0 pts\n"
        "• Confidential Data Copy (USB): -15.0 pts\n"
        "• Large USB Data Transfer (>100MB): -25.0 pts\n"
        "• Webmail / Cloud Upload Navigations: -20.0 pts\n"
        "• Unauthorized Privilege Escalation: -20.0 pts"
    )
    
    doc.add_heading("3.2 Random Forest Classifier Model", level=2)
    doc.add_paragraph(
        "A Random Forest model is trained on a 17-dimensional feature matrix extracted for the 150 "
        "ground-truth employees. Features include activity frequency aggregates (logon counts, device counts, "
        "file accesses, email metrics, web categories) combined with the user's LDAP-mapped "
        "Big Five Psychometric scores (Openness, Conscientiousness, Extraversion, Agreeableness, Neuroticism)."
    )
    
    # Save file
    filename = "GarudaAI_Tech_Stack_and_Model_Guide.docx"
    doc.save(filename)
    print(f"Document created: {filename}")

if __name__ == "__main__":
    create_doc()
